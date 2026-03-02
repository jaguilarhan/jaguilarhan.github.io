"""
Generador de Reporte Word para Cuaderno de Obra
Lee actividades.json exportado desde la app y genera el .docx
Uso: python generar_reporte.py [desde] [hasta] [residente] [inspector]
Ejemplo: python generar_reporte.py 2026-02-17 2026-02-24 "Juan Perez" "Maria Lopez"
"""
import json, sys, os, io
from datetime import datetime, date

# Forzar UTF-8 en stdout/stderr para evitar errores cp1252 en Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Secuencias de proceso (igual que en la app) ──────────────────
SECUENCIAS = {
    "COLUMNAS":        ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "COLUMNETAS":      ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "VIGAS":           ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "VIGUETAS":        ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "PLACAS":          ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "MUROS":           ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "ESCALERAS":       ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "PARAPETOS":       ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "MESONES":         ["Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "ZAPATAS":         ["Excavacion","Solado","Habilitado de Acero","Armado de Acero","Vaciado de Concreto","Curado","Relleno"],
    "CIMIENTOS":       ["Excavacion","Habilitado de Acero","Armado de Acero","Encofrado","Vaciado de Concreto","Desencofrado","Curado"],
    "LOSAS_ALIGERADAS":["Encofrado","Colocacion Ladrillo","Habilitado de Acero","Armado de Acero","Vaciado de Concreto","Desencofrado","Curado"],
    "LOSAS_MACIZAS":   ["Encofrado","Habilitado de Acero","Armado de Acero","Vaciado de Concreto","Desencofrado","Curado"],
}

PASO_CLAVES = {
    "Habilitado de Acero": "ACERO",
    "Armado de Acero":     "ACERO",
    "Encofrado":           "ENCOFRADO",
    "Vaciado de Concreto": "CONCRETO",
    "Desencofrado":        "ENCOFRADO",
    "Curado":              "CURADO",
    "Excavacion":          "EXCAVAC",
    "Solado":              "SOLADO",
    "Relleno":             "RELLENO",
    "Colocacion Ladrillo": "LADRILLO",
}

def calcular_avance(categoria, acts):
    pasos = SECUENCIAS.get(categoria)
    if not pasos:
        return None
    nombres_partidas = [a["partida"]["nombre"].upper() for a in acts]
    hechos = []
    for paso in pasos:
        clave = PASO_CLAVES.get(paso, paso.upper())
        hecho = any(clave in n for n in nombres_partidas)
        hechos.append(hecho)
    total_hechos = sum(hechos)
    pct = round(total_hechos / len(pasos) * 100)
    prox_idx = next((i for i,h in enumerate(hechos) if not h), -1)
    prox_paso = pasos[prox_idx] if prox_idx >= 0 else "COMPLETADO"
    return {"pasos": pasos, "hechos": hechos, "pct": pct,
            "total_hechos": total_hechos, "total": len(pasos),
            "prox_paso": prox_paso}

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_header_row(table, headers, bg="1A237E"):
    row = table.add_row()
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
    return row

def fmt_fecha(fecha_str):
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    dias_sem = ["lunes","martes","miércoles","jueves","viernes","sábado","domingo"]
    d = datetime.strptime(fecha_str, "%Y-%m-%d")
    return f"{dias_sem[d.weekday()].capitalize()} {d.day:02d} de {meses[d.month-1]} de {d.year}"

# ── Leer argumentos ──────────────────────────────────────────────
base = r"C:\Users\Jose Alonso\IntelliJ_IDEA\work_manage_app"
acts_path = os.path.join(base, "actividades_exportadas.json")

args = sys.argv[1:]
desde_str  = args[0] if len(args) > 0 else None
hasta_str  = args[1] if len(args) > 1 else None
residente  = args[2] if len(args) > 2 else "Residente de Obra"
inspector  = args[3] if len(args) > 3 else "Inspector / Supervisor"

# ── Leer actividades ─────────────────────────────────────────────
if not os.path.exists(acts_path):
    print(f"ERROR: No se encontró {acts_path}")
    print("Primero exporta las actividades desde la app (botón Exportar JSON)")
    sys.exit(1)

with open(acts_path, encoding="utf-8") as f:
    todas = json.load(f)

# Filtrar por fechas
if desde_str and hasta_str:
    desde_d = datetime.strptime(desde_str, "%Y-%m-%d").date()
    hasta_d = datetime.strptime(hasta_str, "%Y-%m-%d").date()
    actividades = [a for a in todas
                   if desde_d <= datetime.strptime(a["fecha"],"%Y-%m-%d").date() <= hasta_d]
else:
    actividades = todas
    todas_fechas = sorted(set(a["fecha"] for a in todas))
    desde_str = todas_fechas[0] if todas_fechas else "N/A"
    hasta_str = todas_fechas[-1] if todas_fechas else "N/A"

if not actividades:
    print("No hay actividades en el rango indicado.")
    sys.exit(0)

print(f"Generando reporte para {len(actividades)} actividades ({desde_str} al {hasta_str})...")

# ── Crear documento ───────────────────────────────────────────────
doc = Document()

# Márgenes y orientación horizontal
from docx.enum.section import WD_ORIENT
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width  = new_width
    section.page_height = new_height
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

# ── TÍTULO ────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("CUADERNO DE OBRA")
run.bold = True; run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("CONTROL DE ACTIVIDADES DIARIAS – SUB PRESUPUESTO 01: ESTRUCTURAS")
r2.bold = True; r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("MEJORAMIENTO Y AMPLIACIÓN DE LOS SERVICIOS OPERATIVOS")
r3.font.size = Pt(9); r3.font.color.rgb = RGBColor(0x55,0x55,0x55)
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("LABORATORIO AMBIENTAL SAN AGUSTÍN DE TORATA - PROVINCIA MARISCAL NIETO - MOQUEGUA")
r4.font.size = Pt(9); r4.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_paragraph()

# ── DATOS GENERALES ───────────────────────────────────────────────
tbl_info = doc.add_table(rows=0, cols=4)
tbl_info.style = "Table Grid"

r1 = tbl_info.add_row()
r1.cells[0].text = "Residente de Obra:"
r1.cells[0].paragraphs[0].runs[0].bold = True
r1.cells[1].text = residente
r1.cells[2].text = "Inspector / Supervisor:"
r1.cells[2].paragraphs[0].runs[0].bold = True
r1.cells[3].text = inspector

r2 = tbl_info.add_row()
r2.cells[0].text = "Período del reporte:"
r2.cells[0].paragraphs[0].runs[0].bold = True
fecha_desde_fmt = fmt_fecha(desde_str) if desde_str != "N/A" else desde_str
fecha_hasta_fmt = fmt_fecha(hasta_str) if hasta_str != "N/A" else hasta_str
merged = r2.cells[1].merge(r2.cells[2]).merge(r2.cells[3])
merged.text = f"{fecha_desde_fmt}  al  {fecha_hasta_fmt}"

for row in tbl_info.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            para.runs[0].font.size = Pt(9) if para.runs else None

doc.add_paragraph()

# ── ACTIVIDADES POR DÍA ───────────────────────────────────────────
por_fecha = {}
for a in actividades:
    por_fecha.setdefault(a["fecha"], []).append(a)

for fecha in sorted(por_fecha.keys()):
    acts_dia = por_fecha[fecha]

    # Cabecera del día
    p_dia = doc.add_paragraph()
    p_dia.paragraph_format.space_before = Pt(10)
    p_dia.paragraph_format.space_after  = Pt(2)
    run_dia = p_dia.add_run(f"  {fmt_fecha(fecha).upper()}  ")
    run_dia.bold = True
    run_dia.font.size = Pt(10)
    run_dia.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Fondo azul en el párrafo
    pPr = p_dia._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A237E')
    pPr.append(shd)

    # Tabla de actividades
    tbl = doc.add_table(rows=0, cols=8)
    tbl.style = "Table Grid"

    # Header
    hrow = tbl.add_row()
    cabeceras = ["N\u00b0", "Descripci\u00f3n de Actividad", "C\u00f3digo Partida", "Nombre de Partida", "Ejes", "Elemento", "Nivel", "% Avance"]
    for i, cab in enumerate(cabeceras):
        c = hrow.cells[i]
        c.text = cab
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, "283593")

    # Filas de datos
    for idx, act in enumerate(acts_dia):
        drow = tbl.add_row()
        drow.cells[0].text = str(idx + 1)
        drow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        drow.cells[1].text = act["descripcion"]
        drow.cells[2].text = act["partida"]["codigo"]
        drow.cells[2].paragraphs[0].runs[0].font.name = "Courier New"

        # Celda partida
        p_nom = drow.cells[3].paragraphs[0]
        p_nom.add_run(act["partida"]["nombre"]).font.size = Pt(8)

        # Ejes
        ejes_val = act.get("ejes", "")
        drow.cells[4].text = ejes_val if ejes_val and ejes_val != "\u2014" else ""

        drow.cells[5].text = act["elemento"] if act["elemento"] != "\u2014" else ""
        if act["elemento"] != "\u2014" and drow.cells[5].paragraphs[0].runs:
            drow.cells[5].paragraphs[0].runs[0].bold = True
        drow.cells[6].text = act["nivel"] if act["nivel"] != "\u2014" else ""

        # Porcentaje manual del usuario
        pct_val = act.get("pct", "")
        pct_text = f"{pct_val}%" if pct_val is not None and pct_val != "" else ""
        drow.cells[7].text = pct_text
        if drow.cells[7].paragraphs[0].runs:
            drow.cells[7].paragraphs[0].runs[0].bold = True
        drow.cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, cell in enumerate(drow.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(8.5)
            if i in (0, 2, 4, 5, 6, 7):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Color de fila alternado
        if idx % 2 == 1:
            for cell in drow.cells:
                set_cell_bg(cell, "F5F5F5")

    # Anchos de columnas
    set_col_widths(tbl, [0.6, 3.5, 2.5, 4.0, 2.0, 1.5, 1.5, 1.2])
    doc.add_paragraph()

# ── RESUMEN DE AVANCE POR ELEMENTO ────────────────────────────────
doc.add_page_break()

p_res = doc.add_paragraph()
r_res = p_res.add_run("RESUMEN DE AVANCE POR ELEMENTO ESTRUCTURAL")
r_res.bold = True; r_res.font.size = Pt(13)
r_res.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
p_res.paragraph_format.space_after = Pt(8)

# Agrupar por categoria + elemento + nivel usando TODAS las actividades
grupos = {}
for a in todas:
    k = f'{a["partida"]["categoria"]}||{a["elemento"]}||{a["nivel"]}'
    grupos.setdefault(k, {"categoria":a["partida"]["categoria"],
                          "elemento":a["elemento"],"nivel":a["nivel"],"acts":[]})
    grupos[k]["acts"].append(a)

tbl_av = doc.add_table(rows=0, cols=6)
tbl_av.style = "Table Grid"

hrow_av = tbl_av.add_row()
for i, cab in enumerate(["Elemento","Categor\u00eda","Nivel","Ejes","Pasos Completados","Pr\u00f3ximo Paso"]):
    c = hrow_av.cells[i]
    c.text = cab
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(8.5)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "283593")

for g in sorted(grupos.values(), key=lambda x: x["elemento"]):
    av = calcular_avance(g["categoria"], g["acts"])
    if not av:
        continue
    pasos_hechos = ", ".join(p for p,h in zip(av["pasos"],av["hechos"]) if h) or "(ninguno)"

    # Recopilar ejes de todas las actividades del grupo
    ejes_set = set()
    for a in g["acts"]:
        e = a.get("ejes", "")
        if e and e != "\u2014":
            ejes_set.add(e)
    ejes_txt = ", ".join(sorted(ejes_set)) if ejes_set else ""

    row = tbl_av.add_row()
    row.cells[0].text = g["elemento"]
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = g["categoria"].replace("_"," ")
    row.cells[2].text = g["nivel"]
    row.cells[3].text = ejes_txt

    p_pc = row.cells[4].paragraphs[0]
    p_pc.add_run(f'{av["total_hechos"]}/{av["total"]} pasos ({av["pct"]}%)').bold = True
    p_pc.runs[0].font.size = Pt(8.5)
    p_det = row.cells[4].add_paragraph()
    run_det = p_det.add_run(pasos_hechos)
    run_det.font.size = Pt(7.5)
    run_det.italic = True
    run_det.font.color.rgb = RGBColor(0x55,0x55,0x55)

    prox_run = row.cells[5].paragraphs[0].add_run(av["prox_paso"])
    prox_run.font.size = Pt(8.5)
    prox_run.bold = av["pct"] == 100
    if av["pct"] == 100:
        prox_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    else:
        prox_run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    for i, cell in enumerate(row.cells):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0,2,3) else WD_ALIGN_PARAGRAPH.LEFT

set_col_widths(tbl_av, [2.0, 3.0, 1.5, 2.5, 5.0, 3.5])

# ── TABLA DE PARTIDAS NUEVAS (no aprobadas) ──────────────────────
acts_nuevas = [a for a in actividades if a.get("partida", {}).get("nueva")]
if acts_nuevas:
    doc.add_page_break()

    p_nuevas = doc.add_paragraph()
    r_nuevas = p_nuevas.add_run("ACTIVIDADES CON PARTIDAS NUEVAS (PENDIENTES DE APROBACION)")
    r_nuevas.bold = True; r_nuevas.font.size = Pt(13)
    r_nuevas.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    p_nuevas.paragraph_format.space_after = Pt(4)

    p_nota = doc.add_paragraph()
    r_nota = p_nota.add_run("Nota: Estas partidas corresponden al Expediente Adicional y aun no estan aprobadas para valorizacion. Se registran como referencia para el cuaderno de obra.")
    r_nota.font.size = Pt(8.5)
    r_nota.italic = True
    r_nota.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p_nota.paragraph_format.space_after = Pt(8)

    tbl_nv = doc.add_table(rows=0, cols=7)
    tbl_nv.style = "Table Grid"

    hrow_nv = tbl_nv.add_row()
    for i, cab in enumerate(["N\u00b0", "Fecha", "Descripci\u00f3n", "C\u00f3digo Partida", "Nombre de Partida", "Ejes", "% Avance"]):
        c = hrow_nv.cells[i]
        c.text = cab
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, "C62828")

    for idx, act in enumerate(sorted(acts_nuevas, key=lambda a: a["fecha"])):
        drow = tbl_nv.add_row()
        drow.cells[0].text = str(idx + 1)
        drow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        drow.cells[1].text = act["fecha"]
        drow.cells[2].text = act["descripcion"]
        drow.cells[3].text = act["partida"]["codigo"]
        drow.cells[4].text = act["partida"]["nombre"]
        ejes_val = act.get("ejes", "")
        drow.cells[5].text = ejes_val if ejes_val and ejes_val != "\u2014" else ""
        pct_val = act.get("pct", "")
        drow.cells[6].text = f"{pct_val}%" if pct_val is not None and pct_val != "" else ""
        drow.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for cell in drow.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(8.5)

        if idx % 2 == 1:
            for cell in drow.cells:
                set_cell_bg(cell, "FFF3F3")

    set_col_widths(tbl_nv, [0.6, 2.0, 4.0, 2.5, 4.5, 2.0, 1.2])

# ── FIRMAS ────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

tbl_firma = doc.add_table(rows=1, cols=2)
cell_res = tbl_firma.rows[0].cells[0]
cell_ins = tbl_firma.rows[0].cells[1]

for cell, nombre, cargo in [(cell_res, residente, "Residente de Obra"),
                             (cell_ins, inspector, "Inspector / Supervisor")]:
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("_" * 35).font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(nombre)
    r.bold = True; r.font.size = Pt(10)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(cargo).font.size = Pt(9)
    # Sin bordes
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top','bottom','left','right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tcPr.append(b)

# ── Guardar ───────────────────────────────────────────────────────
nombre_archivo = f"Cuaderno_Obra_{desde_str}_al_{hasta_str}.docx"
out_path = os.path.join(base, nombre_archivo)
doc.save(out_path)
print(f"\n[OK] Reporte generado exitosamente:")
print(f"   {out_path}")

